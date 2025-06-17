from pathlib import Path
from bs4 import BeautifulSoup
import docx
import pandas as pd
from PyPDF2 import PdfReader
import re
from pdf2docx import Converter
from rag.trunk.pdf_convert import extract_pdf_text

def convert_docx_to_markdown(filepath):
    doc = docx.Document(filepath)
    markdown_lines = []
    has_heading = False
    para_index = 0

    for element in doc.element.body:
        if element.tag.endswith('tbl'):  # 表格处理
            table = []
            for row in element.findall('.//w:tr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                cells = []
                for cell in row.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    cells.append(cell.text)
                if cells:
                    table.append('|' + '|'.join(cells) + '|')
            
            if table:
                header_separator = '|' + '|'.join(['---'] * len(table[0].split('|')[1:-1])) + '|'
                table.insert(1, header_separator)
                markdown_lines.extend(table)
                markdown_lines.append('')  # 添加空行
        
        elif element.tag.endswith('p'):  # 段落处理
            if para_index >= len(doc.paragraphs):
                continue
                
            para = doc.paragraphs[para_index]
            text = para.text.strip()
            para_index += 1  # 增加段落索引
            
            if not text:
                continue

            style = para.style.name.lower()

            punctuation = '。！？，；,.!?;'  # 定义标点符号
            if "heading" in style:
                level = 1  # 所有标题都处理为一级标题
                markdown_lines.append(f"# {text}")
                has_heading = True
            elif len(text) <= 20 and text[-1] not in punctuation:  # 新增判断条件
                markdown_lines.append(f"# {text}")
                has_heading = True
            elif "list" in style or para._element.xpath(".//w:numPr"):
                markdown_lines.append(f"- {text}")
            else:
                markdown_lines.append(text)

    if not has_heading and markdown_lines:
        markdown_lines.insert(0, "# 第一段")

    return '\n\n'.join(markdown_lines)

def convert_pdf_to_markdown(pdf_path):
    markdown_text = extract_pdf_text(pdf_path)
    if markdown_text and "#" not in markdown_text:
        markdown_text.insert(0, "# 第一段")
    return markdown_text

def convert_txt_to_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    if paragraphs:
        paragraphs.insert(0, "# 第一段")
    return '\n\n'.join(paragraphs)

def convert_html_to_markdown(filepath):
    """
    将HTML文件转为Markdown字符串
    - 有标题，加#标题
    - 有段落，按段落分
    - 无段落时，加整体段
    - 列表转成普通文本
    - 表格转成Markdown表格
    """
    # 读取文件
    with open(filepath, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    markdown_lines = []
    has_title = False
    has_paragraph = False
    collected_texts = []

    def parse_element(element):
        nonlocal has_title, has_paragraph

        if element.name and element.name.startswith('h'):
            # 标题
            level = int(element.name[1])
            text = element.get_text(strip=True)
            if text:
                markdown_lines.append('#' * level + ' ' + text)
                has_title = True
        elif element.name == 'p':
            # 段落
            text = element.get_text(strip=True)
            if text:
                markdown_lines.append(text)
                has_paragraph = True
        elif element.name in ['ul', 'ol']:
            # 列表，转成普通文本
            for li in element.find_all('li'):
                text = li.get_text(strip=True)
                if text:
                    markdown_lines.append(text)
        elif element.name == 'table':
            # 表格，转成Markdown表格
            rows = element.find_all('tr')
            table_data = []
            for row in rows:
                cols = row.find_all(['td', 'th'])
                table_data.append([col.get_text(strip=True) for col in cols])

            if table_data:
                header = table_data[0]
                markdown_lines.append('| ' + ' | '.join(header) + ' |')
                markdown_lines.append('|' + '|'.join([' --- ' for _ in header]) + '|')
                for row in table_data[1:]:
                    markdown_lines.append('| ' + ' | '.join(row) + ' |')
        elif element.name is not None:
            # 其他元素，收集纯文本备用（如果最终没有段落时用）
            text = element.get_text(strip=True)
            if text:
                collected_texts.append(text)

    # 只处理body
    body = soup.body or soup
    for element in body.find_all(recursive=False):
        parse_element(element)

    # 如果没有标题，加默认标题
    if not has_title:
        markdown_lines.insert(0, '# 第一章')

    # 如果没有段落，整理备用文本为一个整体段
    if not has_paragraph:
        if collected_texts:
            merged_text = ' '.join(collected_texts)
            markdown_lines.append(merged_text)

    markdown_content = '\n\n'.join(markdown_lines)
    return markdown_content

def convert_xlsx_to_markdown(filepath):
    dfs = pd.read_excel(filepath, sheet_name=None)
    markdown_lines = []

    for sheet_name, df in dfs.items():
        markdown_lines.append(f"# {sheet_name}")
        
        # 处理每个表格
        if not df.empty:
            # 获取表头
            headers = df.columns.tolist()
            # 创建表格行
            table = ['|' + '|'.join(str(col) for col in headers) + '|']
            # 添加分隔行
            table.append('|' + '|'.join(['---'] * len(headers)) + '|')
            # 添加数据行
            for _, row in df.iterrows():
                table.append('|' + '|'.join(str(cell) for cell in row) + '|')
            # 将表格添加到markdown_lines
            markdown_lines.extend(table)
            markdown_lines.append('')  # 添加空行

    return '\n\n'.join(markdown_lines)

def convert_csv_to_markdown(filepath):
    """
    将CSV文件转换为Markdown格式
    """
    # 读取CSV文件
    df = pd.read_csv(filepath, encoding='utf-8')
    markdown_lines = []
    
    # 使用文件名作为标题
    filename = Path(filepath).stem
    markdown_lines.append(f"# {filename}")
    
    # 处理表格数据
    if not df.empty:
        # 获取表头
        headers = df.columns.tolist()
        # 创建表格行
        table = ['|' + '|'.join(str(col) for col in headers) + '|']
        # 添加分隔行
        table.append('|' + '|'.join(['---'] * len(headers)) + '|')
        # 添加数据行
        for _, row in df.iterrows():
            # 处理空值，转换为空字符串
            row_data = [str(cell) if pd.notna(cell) else '' for cell in row]
            table.append('|' + '|'.join(row_data) + '|')
        # 将表格添加到markdown_lines
        markdown_lines.extend(table)
    
    return '\n\n'.join(markdown_lines)

def convert_to_markdown(filepath, filetype):
    # 获取不带扩展名的文件名
    filename = Path(filepath).stem
    
    # 根据文件类型进行转换
    if filetype == 'docx':
        content = convert_docx_to_markdown(filepath)
    elif filetype == 'pdf':
        content = convert_pdf_to_markdown(filepath)
    elif filetype == 'txt':
        content = convert_txt_to_markdown(filepath)
    elif filetype == 'html':
        content = convert_html_to_markdown(filepath)
    elif filetype == 'xlsx':
        content = convert_xlsx_to_markdown(filepath)
    elif filetype == 'csv':
        content = convert_csv_to_markdown(filepath)
    else:
        raise ValueError(f"Unsupported file type: {filetype}")
    
    return [filename, content]

# 示例调用
if __name__ == "__main__":
    markdown_text = convert_to_markdown("rag/test_docs/AI信息化在华通公司的实施方案v1.1.pdf", "pdf")
    print(markdown_text[1])